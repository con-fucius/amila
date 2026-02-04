"""
Report Generation Service
Generates professional executive reports from query results
"""

import logging
from typing import Dict, Any, List, Optional
from datetime import datetime, timezone
from io import BytesIO
import json

logger = logging.getLogger(__name__)

# Feature flag
REPORT_GENERATION_ENABLED = True


class ReportGenerationService:
    """
    Service for generating executive reports from query results
    
    Supports formats: PDF, DOCX, HTML
    Template structure:
    - Heading
    - Executive Summary
    - Metrics and Inference
    - Insights/Recommendations
    """
    
    TEMPLATE_SECTIONS = [
        "heading",
        "executive_summary",
        "metrics_inference",
        "insights_recommendations",
        "narrative_summary",
    ]
    
    @classmethod
    async def generate_report(
        cls,
        query_results: List[Dict[str, Any]],
        format: str = "html",
        title: Optional[str] = None,
        user_queries: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        Generate an executive report from query results
        
        Args:
            query_results: List of query result dicts with columns, rows, sql_query
            format: Output format (html, pdf, docx)
            title: Report title
            user_queries: Original user queries for context
            
        Returns:
            Dict with report content and metadata
        """
        if not REPORT_GENERATION_ENABLED:
            return {"status": "error", "message": "Report generation is disabled"}
        
        if not query_results:
            return {"status": "error", "message": "No query results provided"}
        
        try:
            # Generate report content
            report_data = cls._build_report_data(query_results, title, user_queries)
            
            # Generate in requested format
            if format.lower() == "html":
                content = cls._generate_html(report_data)
                content_type = "text/html"
            elif format.lower() == "pdf":
                content = await cls._generate_pdf(report_data)
                content_type = "application/pdf"
            elif format.lower() == "docx":
                content = await cls._generate_docx(report_data)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                return {"status": "error", "message": f"Unsupported format: {format}"}
            
            return {
                "status": "success",
                "format": format,
                "content_type": content_type,
                "content": content,
                "title": report_data["title"],
                "generated_at": datetime.now(timezone.utc).isoformat(),
            }
            
        except Exception as e:
            logger.error(f"Report generation failed: {e}", exc_info=True)
            return {"status": "error", "message": str(e)}
    
    @classmethod
    def _build_report_data(
        cls,
        query_results: List[Dict[str, Any]],
        title: Optional[str],
        user_queries: Optional[List[str]],
    ) -> Dict[str, Any]:
        """Build structured report data from query results"""
        
        # Generate title if not provided
        if not title:
            title = f"Data Analysis Report - {datetime.now().strftime('%B %d, %Y')}"
        
        # Analyze results for insights
        total_rows = sum(r.get("row_count", len(r.get("rows", []))) for r in query_results)
        total_queries = len(query_results)
        
        # Extract key metrics from results
        metrics = []
        for i, result in enumerate(query_results):
            columns = result.get("columns", [])
            rows = result.get("rows", [])
            
            if not rows:
                continue
            
            def get_val(r, idx, name):
                if isinstance(r, dict):
                    return r.get(name)
                try:
                    return r[idx]
                except (IndexError, TypeError):
                    return None

            # Find numeric columns for metrics
            for col_idx, col in enumerate(columns):
                values = [get_val(row, col_idx, col) for row in rows if get_val(row, col_idx, col) is not None]
                numeric_values = []
                for v in values:
                    try:
                        numeric_values.append(float(v))
                    except (ValueError, TypeError):
                        pass
                
                if numeric_values:
                    metrics.append({
                        "name": col,
                        "query_index": i,
                        "count": len(numeric_values),
                        "sum": sum(numeric_values),
                        "avg": sum(numeric_values) / len(numeric_values),
                        "min": min(numeric_values),
                        "max": max(numeric_values),
                    })
        
        # Build executive summary
        summary_points = [
            f"This report analyzes {total_queries} data queries with {total_rows:,} total records.",
        ]
        
        if metrics:
            top_metric = max(metrics, key=lambda m: m["sum"])
            summary_points.append(
                f"Key metric: {top_metric['name']} with total of {top_metric['sum']:,.2f} "
                f"(avg: {top_metric['avg']:,.2f})"
            )
        
        # Generate insights (rule-based for sync context)
        # LLM insights are generated separately in async context
        insights = cls._generate_insights(query_results, metrics)
        
        return {
            "title": title,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "executive_summary": summary_points,
            "query_count": total_queries,
            "total_rows": total_rows,
            "metrics": metrics[:10],  # Top 10 metrics
            "insights": insights,
            "narrative": "",
            "query_results": query_results,
            "user_queries": user_queries or [],
        }
    
    @classmethod
    async def generate_report_with_llm_insights(
        cls,
        query_results: List[Dict[str, Any]],
        format: str = "html",
        title: Optional[str] = None,
        user_queries: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        Generate report with LLM-enhanced insights.
        
        This is the preferred method when async context is available.
        Uses anti-hallucination constraints to ensure insights are grounded in data.
        """
        if not REPORT_GENERATION_ENABLED:
            return {"status": "error", "message": "Report generation is disabled"}
        
        if not query_results:
            return {"status": "error", "message": "No query results provided"}
        
        try:
            # Build base report data
            report_data = cls._build_report_data(query_results, title, user_queries)
            
            # Extract metrics for LLM insights
            metrics = report_data.get("metrics", [])
            
            # Generate LLM-enhanced insights with anti-hallucination
            llm_insights = await cls._generate_llm_insights(
                query_results, 
                metrics, 
                user_queries
            )
            
            # Replace rule-based insights with LLM insights
            report_data["insights"] = llm_insights
            report_data["insights_source"] = "llm_enhanced"

            # Narrative storytelling
            try:
                from app.services.narrative_service import NarrativeService
                report_data["narrative"] = await NarrativeService.generate_narrative(
                    query_results=query_results,
                    metrics=metrics,
                    trends=[],
                    anomalies=[],
                    user_queries=user_queries
                )
            except Exception:
                report_data["narrative"] = ""
            
            # Generate in requested format
            if format.lower() == "html":
                content = cls._generate_html(report_data)
                content_type = "text/html"
            elif format.lower() == "pdf":
                content = await cls._generate_pdf(report_data)
                content_type = "application/pdf"
            elif format.lower() == "docx":
                content = await cls._generate_docx(report_data)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                return {"status": "error", "message": f"Unsupported format: {format}"}
            
            return {
                "status": "success",
                "format": format,
                "content_type": content_type,
                "content": content,
                "title": report_data["title"],
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "insights_source": "llm_enhanced",
            }
            
        except Exception as e:
            logger.error(f"Report generation with LLM insights failed: {e}", exc_info=True)
            # Fallback to rule-based insights
            return await cls.generate_report(query_results, format, title, user_queries)

    @classmethod
    def _generate_insights(
        cls,
        query_results: List[Dict[str, Any]],
        metrics: List[Dict[str, Any]],
    ) -> List[str]:
        """Generate data-driven insights from results (rule-based)"""
        insights = []
        
        # Analyze trends in metrics
        for metric in metrics[:5]:
            if metric["max"] > metric["avg"] * 2:
                insights.append(
                    f"Notable outlier detected in {metric['name']}: "
                    f"maximum value ({metric['max']:,.2f}) is significantly higher than average ({metric['avg']:,.2f})"
                )
            
            if metric["min"] < metric["avg"] * 0.1 and metric["min"] >= 0:
                insights.append(
                    f"Wide variance in {metric['name']}: "
                    f"minimum ({metric['min']:,.2f}) is much lower than average ({metric['avg']:,.2f})"
                )
        
        # Analyze result patterns
        for i, result in enumerate(query_results):
            rows = result.get("rows", [])
            if len(rows) == 0:
                insights.append(f"Query {i+1} returned no results - consider broadening search criteria")
            elif len(rows) >= 1000:
                insights.append(f"Query {i+1} returned {len(rows):,} rows - consider adding filters for focused analysis")
        
        if not insights:
            insights.append("Data appears consistent with no significant anomalies detected")
        
        return insights
    
    @classmethod
    async def _generate_llm_insights(
        cls,
        query_results: List[Dict[str, Any]],
        metrics: List[Dict[str, Any]],
        user_queries: Optional[List[str]] = None,
    ) -> List[str]:
        """
        Generate LLM-enhanced insights with anti-hallucination constraints
        
        Uses constrained prompts to ensure insights are grounded in actual data.
        """
        try:
            from app.orchestrator.llm_config import get_llm
            
            llm = get_llm()
            if not llm:
                logger.warning("LLM not available for insights, using rule-based")
                return cls._generate_insights(query_results, metrics)
            
            # Build data summary for LLM
            data_summary = []
            for i, result in enumerate(query_results[:3]):
                columns = result.get("columns", [])
                rows = result.get("rows", [])[:5]  # First 5 rows only
                row_count = result.get("row_count", len(result.get("rows", [])))
                
                query_text = user_queries[i] if user_queries and i < len(user_queries) else f"Query {i+1}"
                
                data_summary.append(f"Query: {query_text}")
                data_summary.append(f"Columns: {', '.join(columns[:8])}")
                data_summary.append(f"Total rows: {row_count}")
                if rows:
                    data_summary.append(f"Sample data: {json.dumps(rows[:3], default=str)[:500]}")
                data_summary.append("")
            
            metrics_summary = []
            for m in metrics[:5]:
                metrics_summary.append(
                    f"- {m['name']}: sum={m['sum']:,.2f}, avg={m['avg']:,.2f}, min={m['min']:,.2f}, max={m['max']:,.2f}"
                )
            
            # Constrained prompt to prevent hallucination
            prompt = f"""Analyze the following data and generate 3-5 business insights.

CRITICAL RULES:
1. ONLY reference values that appear in the data below
2. DO NOT speculate or make assumptions beyond the data
3. Each insight MUST cite specific values from the data
4. If data is insufficient, say so rather than guessing

DATA:
{chr(10).join(data_summary)}

METRICS:
{chr(10).join(metrics_summary)}

Generate insights in this format:
- [Insight with specific data citation]

Example good insight: "Revenue of $456,789 from Widget A represents 37% of total, making it the top performer."
Example bad insight: "Sales are likely to increase next quarter." (speculation - DO NOT do this)

Your insights:"""
            
            response = await llm.ainvoke(prompt)
            response_text = response.content if hasattr(response, 'content') else str(response)
            
            # Parse insights from response
            llm_insights = []
            for line in response_text.strip().split('\n'):
                line = line.strip()
                if line.startswith('-') or line.startswith('*'):
                    insight = line.lstrip('-*').strip()
                    if insight and len(insight) > 10:
                        llm_insights.append(insight)
            
            # Combine with rule-based insights
            rule_insights = cls._generate_insights(query_results, metrics)
            
            # Deduplicate and limit
            all_insights = llm_insights[:3] + rule_insights[:2]
            seen = set()
            unique_insights = []
            for insight in all_insights:
                key = insight[:50].lower()
                if key not in seen:
                    seen.add(key)
                    unique_insights.append(insight)
            
            return unique_insights[:5]
            
        except Exception as e:
            logger.warning(f"LLM insights generation failed: {e}")
            return cls._generate_insights(query_results, metrics)
    
    @classmethod
    def _generate_html(cls, report_data: Dict[str, Any]) -> str:
        """Generate HTML report"""
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report_data['title']}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #333; max-width: 900px; margin: 0 auto; padding: 40px 20px; }}
        .header {{ text-align: center; margin-bottom: 40px; padding-bottom: 20px; border-bottom: 3px solid #10b981; }}
        .header h1 {{ color: #111827; font-size: 28px; margin-bottom: 10px; }}
        .header .date {{ color: #6b7280; font-size: 14px; }}
        .section {{ margin-bottom: 30px; }}
        .section h2 {{ color: #10b981; font-size: 20px; margin-bottom: 15px; padding-bottom: 8px; border-bottom: 1px solid #e5e7eb; }}
        .summary-list {{ list-style: none; }}
        .summary-list li {{ padding: 8px 0; padding-left: 20px; position: relative; }}
        .summary-list li::before {{ content: "-"; color: #10b981; position: absolute; left: 0; font-weight: bold; }}
        .metrics-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; }}
        .metric-card {{ background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 8px; padding: 15px; }}
        .metric-card .name {{ font-weight: 600; color: #374151; margin-bottom: 8px; }}
        .metric-card .value {{ font-size: 24px; color: #10b981; font-weight: bold; }}
        .metric-card .details {{ font-size: 12px; color: #6b7280; margin-top: 5px; }}
        .insights-list {{ background: #f0fdf4; border-left: 4px solid #10b981; padding: 15px 20px; }}
        .insights-list li {{ margin-bottom: 10px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 13px; }}
        th, td {{ padding: 10px; text-align: left; border-bottom: 1px solid #e5e7eb; }}
        th {{ background: #f9fafb; font-weight: 600; color: #374151; }}
        tr:hover {{ background: #f9fafb; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #e5e7eb; text-align: center; color: #9ca3af; font-size: 12px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{report_data['title']}</h1>
        <div class="date">Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</div>
    </div>
    
    <div class="section">
        <h2>Executive Summary</h2>
        <ul class="summary-list">
"""
        for point in report_data.get("executive_summary", []):
            html += f"            <li>{point}</li>\n"
        
        html += """        </ul>
    </div>
    
    <div class="section">
        <h2>Key Metrics</h2>
        <div class="metrics-grid">
"""
        for metric in report_data.get("metrics", [])[:6]:
            html += f"""            <div class="metric-card">
                <div class="name">{metric['name']}</div>
                <div class="value">{metric['sum']:,.2f}</div>
                <div class="details">Avg: {metric['avg']:,.2f} | Min: {metric['min']:,.2f} | Max: {metric['max']:,.2f}</div>
            </div>
"""
        
        html += """        </div>
    </div>
    
    <div class="section">
        <h2>Insights & Recommendations</h2>
        <ul class="insights-list">
"""
        for insight in report_data.get("insights", []):
            html += f"            <li>{insight}</li>\n"
        
        html += """        </ul>
    </div>
"""

        # Narrative section
        narrative = report_data.get("narrative")
        if narrative:
            html += f"""
    <div class="section">
        <h2>Narrative Summary</h2>
        <div class="insights-list">
            <li>{narrative}</li>
        </div>
    </div>
"""
        
        # Add data tables
        for i, result in enumerate(report_data.get("query_results", [])[:3]):
            columns = result.get("columns", [])
            rows = result.get("rows", [])[:10]  # First 10 rows
            
            if columns and rows:
                query_text = report_data.get("user_queries", [])[i] if i < len(report_data.get("user_queries", [])) else f"Query {i+1}"
                html += f"""
    <div class="section">
        <h2>Data: {query_text[:50]}{'...' if len(query_text) > 50 else ''}</h2>
        <table>
            <thead>
                <tr>
"""
                for col in columns[:8]:  # Max 8 columns
                    html += f"                    <th>{col}</th>\n"
                html += """                </tr>
            </thead>
            <tbody>
"""
                for row in rows:
                    html += "                <tr>\n"
                    for cell in row[:8]:
                        cell_str = str(cell) if cell is not None else "NULL"
                        html += f"                    <td>{cell_str[:50]}</td>\n"
                    html += "                </tr>\n"
                
                html += """            </tbody>
        </table>
    </div>
"""
        
        html += f"""
    <div class="footer">
        <p>Report generated by AMILA Business Intelligence System</p>
        <p>Total queries analyzed: {report_data['query_count']} | Total records: {report_data['total_rows']:,}</p>
    </div>
</body>
</html>"""
        
        return html
    
    @classmethod
    async def _generate_pdf(cls, report_data: Dict[str, Any]) -> bytes:
        """
        Generate PDF report using weasyprint with chart embedding.
        
        Features:
        - Professional HTML-to-PDF rendering
        - Embedded charts (if chart_data provided)
        - Customizable branding via report_data['branding']
        - Table pagination for large result sets
        """
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
            
            # Generate enhanced HTML with PDF-optimized styles
            html_content = cls._generate_pdf_html(report_data)
            
            # Configure fonts
            font_config = FontConfiguration()
            
            # Convert to PDF with custom CSS for print optimization
            css = CSS(string=cls._get_pdf_styles(), font_config=font_config)
            pdf_bytes = HTML(string=html_content).write_pdf(stylesheets=[css])
            
            logger.info(f"Generated PDF report: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except ImportError as e:
            logger.error(f"weasyprint not available: {e}")
            # Fallback: return HTML as bytes with proper headers
            html_content = cls._generate_html(report_data)
            return html_content.encode('utf-8')
        except Exception as e:
            logger.error(f"PDF generation failed: {e}", exc_info=True)
            raise ValueError(f"Failed to generate PDF: {e}")
    
    @classmethod
    def _generate_pdf_html(cls, report_data: Dict[str, Any]) -> str:
        """Generate HTML optimized for PDF rendering with embedded charts"""
        
        branding = report_data.get('branding', {})
        logo_url = branding.get('logo_url', '')
        primary_color = branding.get('primary_color', '#10b981')
        company_name = branding.get('company_name', 'AMILA Business Intelligence')
        
        # Generate chart sections if chart data provided
        chart_sections = ''
        charts = report_data.get('charts', [])
        for i, chart in enumerate(charts[:3]):  # Max 3 charts
            chart_html = cls._generate_chart_embed(chart, i)
            chart_sections += chart_html
        
        # Build data tables with pagination hints
        tables_html = ''
        for i, result in enumerate(report_data.get('query_results', [])[:3]):
            tables_html += cls._generate_pdf_table(result, i, report_data.get('user_queries', []))
        
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report_data['title']}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
            @bottom-center {{
                content: "Page " counter(page) " of " counter(pages);
                font-size: 9pt;
                color: #6b7280;
            }}
            @bottom-right {{
                content: "Generated by {company_name}";
                font-size: 8pt;
                color: #9ca3af;
            }}
        }}
        body {{
            font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            color: #1f2937;
            font-size: 10pt;
        }}
        .header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 3px solid {primary_color};
        }}
        .logo {{
            max-height: 60px;
            margin-bottom: 15px;
        }}
        .header h1 {{
            color: #111827;
            font-size: 22pt;
            margin: 0 0 10px 0;
            font-weight: 600;
        }}
        .header .date {{
            color: #6b7280;
            font-size: 10pt;
        }}
        .section {{
            margin-bottom: 25px;
            page-break-inside: avoid;
        }}
        .section h2 {{
            color: {primary_color};
            font-size: 14pt;
            margin: 0 0 12px 0;
            padding-bottom: 8px;
            border-bottom: 1px solid #e5e7eb;
            font-weight: 600;
        }}
        .summary-list {{
            list-style: none;
            padding: 0;
            margin: 0;
        }}
        .summary-list li {{
            padding: 8px 0;
            padding-left: 20px;
            position: relative;
        }}
        .summary-list li::before {{
            content: "•";
            color: {primary_color};
            position: absolute;
            left: 0;
            font-weight: bold;
            font-size: 14pt;
            line-height: 1.2;
        }}
        .metrics-grid {{
            display: flex;
            flex-wrap: wrap;
            gap: 15px;
        }}
        .metric-card {{
            background: #f9fafb;
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            padding: 15px;
            min-width: 150px;
            flex: 1;
        }}
        .metric-card .name {{
            font-weight: 600;
            color: #374151;
            margin-bottom: 8px;
            font-size: 9pt;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .metric-card .value {{
            font-size: 18pt;
            color: {primary_color};
            font-weight: bold;
        }}
        .metric-card .details {{
            font-size: 8pt;
            color: #6b7280;
            margin-top: 5px;
        }}
        .insights-list {{
            background: #f0fdf4;
            border-left: 4px solid {primary_color};
            padding: 15px 20px;
            margin: 0;
        }}
        .insights-list li {{
            margin-bottom: 8px;
            font-size: 10pt;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
            font-size: 9pt;
            page-break-inside: auto;
        }}
        thead {{
            display: table-header-group;
        }}
        tr {{
            page-break-inside: avoid;
        }}
        th, td {{
            padding: 8px 10px;
            text-align: left;
            border-bottom: 1px solid #e5e7eb;
        }}
        th {{
            background: #f9fafb;
            font-weight: 600;
            color: #374151;
            font-size: 8pt;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        td {{
            font-size: 9pt;
        }}
        .chart-container {{
            margin: 20px 0;
            text-align: center;
            page-break-inside: avoid;
        }}
        .chart-image {{
            max-width: 100%;
            height: auto;
            border: 1px solid #e5e7eb;
            border-radius: 8px;
        }}
        .chart-caption {{
            font-size: 9pt;
            color: #6b7280;
            margin-top: 8px;
            font-style: italic;
        }}
        .query-header {{
            background: #f3f4f6;
            padding: 10px 15px;
            margin: 15px 0 10px 0;
            border-radius: 6px;
            font-weight: 600;
            color: #374151;
            font-size: 10pt;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e5e7eb;
            text-align: center;
            color: #9ca3af;
            font-size: 8pt;
        }}
        .page-break {{
            page-break-before: always;
        }}
    </style>
</head>
<body>
    <div class="header">
        {f'<img src="{logo_url}" class="logo" alt="Company Logo" />' if logo_url else ''}
        <h1>{report_data['title']}</h1>
        <div class="date">Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</div>
    </div>
    
    <div class="section">
        <h2>Executive Summary</h2>
        <ul class="summary-list">
            {''.join(f'<li>{point}</li>' for point in report_data.get('executive_summary', []))}
        </ul>
    </div>
    
    <div class="section">
        <h2>Key Metrics</h2>
        <div class="metrics-grid">
            {''.join(f'''
            <div class="metric-card">
                <div class="name">{metric['name']}</div>
                <div class="value">{metric['sum']:,.2f}</div>
                <div class="details">Avg: {metric['avg']:,.2f} | Min: {metric['min']:,.2f} | Max: {metric['max']:,.2f}</div>
            </div>
            ''' for metric in report_data.get('metrics', [])[:6])}
        </div>
    </div>
    
    <div class="section">
        <h2>Insights & Recommendations</h2>
        <ul class="insights-list">
            {''.join(f'<li>{insight}</li>' for insight in report_data.get('insights', []))}
        </ul>
    </div>

    {f'''
    <div class="section">
        <h2>Narrative Summary</h2>
        <ul class="insights-list">
            <li>{report_data.get('narrative')}</li>
        </ul>
    </div>
    ''' if report_data.get('narrative') else ''}
    
    {chart_sections}
    
    <div class="page-break"></div>
    
    <div class="section">
        <h2>Detailed Data</h2>
        {tables_html}
    </div>
    
    <div class="footer">
        <p>Report generated by {company_name}</p>
        <p>Total queries analyzed: {report_data['query_count']} | Total records: {report_data['total_rows']:,}</p>
        <p>Query ID: {report_data.get('query_id', 'N/A')}</p>
    </div>
</body>
</html>"""
    
    @classmethod
    def _generate_chart_embed(cls, chart: Dict[str, Any], index: int) -> str:
        """Generate HTML for an embedded chart"""
        chart_type = chart.get('type', 'bar')
        chart_title = chart.get('title', f'Chart {index + 1}')
        
        # If chart image data provided (base64), embed directly
        if 'image_data' in chart:
            mime_type = chart.get('mime_type', 'image/png')
            return f'''
            <div class="section chart-container">
                <h2>{chart_title}</h2>
                <img src="data:{mime_type};base64,{chart['image_data']}" class="chart-image" alt="{chart_title}" />
                <div class="chart-caption">{chart.get('caption', '')}</div>
            </div>
            '''
        
        # Otherwise, generate a simple SVG placeholder
        return f'''
        <div class="section chart-container">
            <h2>{chart_title}</h2>
            <div class="chart-caption">Chart data available in interactive version</div>
        </div>
        '''
    
    @classmethod
    def _generate_pdf_table(cls, result: Dict[str, Any], index: int, user_queries: List[str]) -> str:
        """Generate HTML table for PDF with proper pagination"""
        columns = result.get('columns', [])
        rows = result.get('rows', [])[:50]  # Limit to 50 rows for PDF
        row_count = result.get('row_count', len(result.get('rows', [])))
        
        if not columns or not rows:
            return ''
        
        query_text = user_queries[index] if index < len(user_queries) else f'Query {index + 1}'
        
        html = f'''
        <div class="query-header">{query_text[:80]}{'...' if len(query_text) > 80 else ''}</div>
        <table>
            <thead>
                <tr>
                    {''.join(f'<th>{col}</th>' for col in columns[:8])}
                </tr>
            </thead>
            <tbody>
        '''
        
        for row in rows:
            html += '<tr>' + ''.join(f'<td>{cls._format_cell(cell)}</td>' for cell in row[:8]) + '</tr>'
        
        html += f'''
            </tbody>
        </table>
        <div style="font-size: 8pt; color: #6b7280; margin-top: 5px;">
            Showing {len(rows)} of {row_count:,} rows
        </div>
        '''
        
        return html
    
    @classmethod
    def _format_cell(cls, cell: Any) -> str:
        """Format a cell value for PDF display"""
        if cell is None:
            return '<span style="color: #9ca3af; font-style: italic;">NULL</span>'
        if isinstance(cell, (int, float)):
            return f"{cell:,.2f}" if isinstance(cell, float) else f"{cell:,}"
        cell_str = str(cell)
        if len(cell_str) > 50:
            return cell_str[:47] + '...'
        return cell_str
    
    @classmethod
    def _get_pdf_styles(cls) -> str:
        """Get additional CSS styles for PDF optimization"""
        return """
        @media print {
            body {
                print-color-adjust: exact;
                -webkit-print-color-adjust: exact;
            }
        }
        """
    
    @classmethod
    async def _generate_docx(cls, report_data: Dict[str, Any]) -> bytes:
        """Generate DOCX report using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(report_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            doc.add_paragraph()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            for point in report_data.get("executive_summary", []):
                doc.add_paragraph(point, style='List Bullet')
            
            # Key Metrics
            doc.add_heading('Key Metrics', level=1)
            for metric in report_data.get("metrics", [])[:6]:
                p = doc.add_paragraph()
                p.add_run(f"{metric['name']}: ").bold = True
                p.add_run(f"{metric['sum']:,.2f} (Avg: {metric['avg']:,.2f})")
            
            # Insights
            doc.add_heading('Insights & Recommendations', level=1)
            for insight in report_data.get("insights", []):
                doc.add_paragraph(insight, style='List Bullet')

            # Narrative Summary
            if report_data.get("narrative"):
                doc.add_heading('Narrative Summary', level=1)
                doc.add_paragraph(report_data["narrative"])
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
            
        except ImportError:
            logger.warning("python-docx not available")
            raise ValueError("DOCX generation requires python-docx package")
